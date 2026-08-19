# Unit tests for user document extraction (pypdfium2 + rapidocr + docx/txt).
from __future__ import annotations

from pathlib import Path

import pytest

from legalrag.ingest.extract import (
    Extraction,
    PageText,
    extractPageText,
    extractText,
)


class TestExtractionModel:
    def test_toDict_shape(self):
        ex = Extraction(source="doc", path="/x.pdf", pages=[PageText(0, "hi", "text")])
        d = ex.toDict()
        assert d["n_pages"] == 1
        assert d["methods"] == ["text"]
        assert d["pages"][0]["chars"] == 2

    def test_full_text_joins_pages(self):
        ex = Extraction(source="doc", path="/x.pdf", pages=[PageText(0, "a", "text"), PageText(1, "b", "text")])
        assert ex.full_text == "a\n\nb"


class TestExtractText:
    def test_txt_file(self, tmp_path: Path):
        p = tmp_path / "lease.txt"
        p.write_text("First line.\nSecond line.")
        ex = extractText(str(p))
        assert ex.source == "lease"
        assert ex.n_pages == 1
        assert ex.methods == {"raw"}
        assert ex.pages[0].text == "First line.\nSecond line."

    def test_missing_file(self, tmp_path: Path):
        with pytest.raises(FileNotFoundError):
            extractText(str(tmp_path / "nope.pdf"))

    def test_directory(self, tmp_path: Path):
        with pytest.raises(IsADirectoryError):
            extractText(str(tmp_path))

    def test_unsupported_extension(self, tmp_path: Path):
        p = tmp_path / "doc.xyz"
        p.write_text("hi")
        with pytest.raises(ValueError, match="unsupported file type"):
            extractText(str(p))

    def test_empty_txt_is_error(self, tmp_path: Path):
        p = tmp_path / "empty.txt"
        p.write_text("  \n  ")
        ex = extractText(str(p))
        assert ex.methods == {"error"}
        assert ex.pages[0].text.strip() == ""

    def test_docx(self, tmp_path: Path):
        import docx

        doc = docx.Document()
        doc.add_paragraph("Clause one.")
        doc.add_paragraph("Clause two.")
        p = tmp_path / "lease.docx"
        doc.save(str(p))
        ex = extractText(str(p))
        assert ex.methods == {"paragraphs"}
        assert "Clause one." in ex.pages[0].text
        assert "Clause two." in ex.pages[0].text

    def test_docx_tables(self, tmp_path: Path):
        import docx

        doc = docx.Document()
        doc.add_paragraph("Clause one.")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Base Rent"
        table.rows[0].cells[1].text = "$1,200/mo"
        p = tmp_path / "lease.docx"
        doc.save(str(p))
        ex = extractText(str(p))
        assert ex.methods == {"paragraphs"}
        text = ex.pages[0].text
        assert "Clause one." in text
        assert "Base Rent" in text
        assert "$1,200/mo" in text

    def test_extension_case_insensitive(self, tmp_path: Path):
        p = tmp_path / "lease.TXT"
        p.write_text("plain")
        ex = extractText(str(p))
        assert ex.pages[0].text == "plain"


class TestPdfTextLayer:
    def test_text_layer_pdf(self, tmp_path: Path):
        from reportlab.pdfgen import canvas

        out = tmp_path / "text.pdf"
        c = canvas.Canvas(str(out))
        c.drawString(72, 720, "Tenant shall pay rent in accordance with this agreement")
        c.save()

        ex = extractText(str(out))
        assert ex.n_pages == 1
        assert ex.methods == {"text"}
        assert "rent" in ex.pages[0].text

    def test_text_layer_pdf_does_not_load_ocr(self, tmp_path: Path, monkeypatch):
        """Pure text-layer PDFs never instantiate the OCR engine (lazy load)."""
        from reportlab.pdfgen import canvas

        out = tmp_path / "text.pdf"
        c = canvas.Canvas(str(out))
        c.drawString(72, 720, "Tenant shall pay rent in accordance with this agreement")
        c.save()

        def _boom():
            raise AssertionError("OCR engine should not be loaded for text PDFs")

        monkeypatch.setattr("legalrag.ingest.extract._ocr_instance", _boom)
        ex = extractText(str(out))
        assert ex.methods == {"text"}

    def test_one_bad_page_keeps_rest(self, tmp_path: Path, monkeypatch):
        """A page that errors is recorded as 'error'; the rest of the PDF survives."""
        from reportlab.pdfgen import canvas

        out = tmp_path / "two_page.pdf"
        c = canvas.Canvas(str(out))
        c.drawString(72, 720, "Page one has sufficient text content here")
        c.showPage()
        c.drawString(72, 720, "Page two also has sufficient text content here")
        c.save()

        import legalrag.ingest.extract as ex_mod

        real_extract = ex_mod.extractPageText
        calls = {"n": 0}

        def _selective(page):
            calls["n"] += 1
            if calls["n"] == 2:  # second page fails
                return "", "error"
            return real_extract(page)

        monkeypatch.setattr(ex_mod, "extractPageText", _selective)
        ex = extractText(str(out))
        assert ex.n_pages == 2
        assert ex.methods == {"text", "error"}
        assert ex.pages[1].method == "error"
        assert ex.pages[0].method == "text"


class FakeTextPage:
    """Stands in for a pypdfium2 page with a fixed text layer."""

    def __init__(self, text: str):
        self._text = text

    def get_textpage(self):
        return self

    def get_text_bounded(self):
        return self._text

    def render(self, scale=None):
        from PIL import Image

        class _FakeBitmap:
            @staticmethod
            def to_pil():
                return Image.new("RGB", (10, 10), "white")

        return _FakeBitmap()


class TestExtractPageText:
    def test_sufficient_text_layer(self):
        text, method = extractPageText(FakeTextPage("x" * 50))
        assert method == "text"
        assert len(text) >= 40

    def test_short_text_falls_back_to_ocr(self):
        def _fake_engine(_):
            return ([[[0, 0, 10, 10]]], ["OCR'd content"], [0.95])

        text, method = extractPageText(FakeTextPage("short"), ocr_engine=_fake_engine)
        assert method == "ocr"
        assert text == "OCR'd content"

    def test_ocr_output_object_supported(self):
        """Modern rapidocr returns a RapidOCROutput object, not a tuple."""
        from legalrag.ingest.extract import _ocr_text

        class FakeOutput:
            def __init__(self):
                self.txts = ["line one", "line two"]
                self.scores = [0.9, 0.8]

        assert _ocr_text(b"", lambda _: FakeOutput()) == "line one\nline two"

    def test_ocr_output_object_empty(self):
        from legalrag.ingest.extract import _ocr_text

        class FakeOutput:
            def __init__(self):
                self.txts = []
                self.scores = []

        assert _ocr_text(b"", lambda _: FakeOutput()) == ""

    def test_ocr_engine_failure_returns_error(self):
        def _fail(_):
            raise RuntimeError("ocr blew up")

        text, method = extractPageText(FakeTextPage(""), ocr_engine=_fail)
        assert method == "error"
        assert text == ""

    def test_text_layer_failure_returns_error(self):
        class BrokenPage:
            def get_textpage(self):
                raise RuntimeError("bad page")

        text, method = extractPageText(BrokenPage())
        assert method == "error"
        assert text == ""