# User document extraction: file -> page/chunk text records.
#
# PDFs use pypdfium2 for the text layer; pages that yield almost no text are
# treated as scanned and OCR'd lazily with rapidocr (ONNX). Images are OCR'd
# directly, docx via python-docx, and plain text/markdown read as-is. No new
# heavyweight dependencies: everything ships in the existing venv.
#
# The rapidocr engine is only instantiated on first OCR need, so text-layer
# PDFs never pay the model-load cost.
#
# Hardening: extraction is best-effort and page-isolated. A page that fails in
# the text layer or OCR is recorded with an empty "error" record instead of
# aborting the document, so one bad page never loses the whole file.
from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path

# A page with fewer than this many text characters is treated as scanned.
MIN_TEXT_CHARS = 40

# render scale for OCR: 2.0 upscales scanned pages for better detection
OCR_SCALE = 2.0

TEXT_EXTS = {".txt", ".md"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp", ".bmp"}
DOCX_EXTS = {".docx"}


@dataclass
class PageText:
    """Text content of one page/chunk plus how it was obtained.

    method is one of "text" (PDF text layer), "ocr" (rendered + OCR'd),
    "paragraphs" (docx), "raw" (txt/md), or "error" (page could not be read).
    """

    page_idx: int
    text: str
    method: str


@dataclass
class Extraction:
    """Extracted text for a whole document."""

    source: str
    path: str
    pages: list[PageText] = field(default_factory=list)

    @property
    def n_pages(self) -> int:
        return len(self.pages)

    @property
    def methods(self) -> set[str]:
        return {p.method for p in self.pages}

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    def toDict(self) -> dict:
        return {
            "source": self.source,
            "path": self.path,
            "n_pages": self.n_pages,
            "methods": sorted(self.methods),
            "pages": [
                {"page_idx": p.page_idx, "method": p.method, "chars": len(p.text)}
                for p in self.pages
            ],
        }


_OCR_ENGINE = None
_OCR_ENGINE_FAILED = False


def _ocr_instance():
    """Lazily-created, cached RapidOCR engine (loaded only on first OCR need).

    Returns None if the engine cannot be loaded (e.g. onnxruntime missing);
    callers treat a None engine as an unavailable OCR path and degrade to
    "error" pages instead of crashing the whole extraction.
    """
    global _OCR_ENGINE, _OCR_ENGINE_FAILED
    if _OCR_ENGINE is not None or _OCR_ENGINE_FAILED:
        return _OCR_ENGINE
    try:
        from rapidocr import RapidOCR

        _OCR_ENGINE = RapidOCR()
    except Exception:  # noqa: BLE001
        _OCR_ENGINE_FAILED = True
        _OCR_ENGINE = None
    return _OCR_ENGINE


def _ocr_text(image_bytes: bytes, engine) -> str:
    """OCR a rendered page image, joining detected text boxes by line.

    Handles both the legacy tuple result (boxes, texts, scores) and the
    modern rapidocr `RapidOCROutput` object (`.txts`, `.scores`).
    """
    try:
        result = engine(image_bytes)
    except Exception:  # noqa: BLE001
        return ""
    if not result:
        return ""
    if hasattr(result, "txts"):
        texts = result.txts or []
    else:
        try:
            _, texts, _ = result
        except (TypeError, ValueError):
            return ""
    lines = [str(line) for line in texts or []]
    return "\n".join(lines).strip()


def extractPageText(page, ocr_engine=None) -> tuple[str, str]:
    """Text + method for one pypdfium2 page. OCRs if the text layer is empty.

    The OCR engine is only loaded when a page actually needs it (lazy); pure
    text-layer documents never pay the model-load cost. Never raises for a bad
    page: returns ("", "error") on failure so callers keep the rest of the doc.
    """
    try:
        textpage = page.get_textpage()
        text = textpage.get_text_bounded().strip()
    except Exception:  # noqa: BLE001
        text = ""
    if len(text) >= MIN_TEXT_CHARS:
        return text, "text"
    if ocr_engine is None:
        ocr_engine = _ocr_instance()
    if ocr_engine is None:
        # OCR unavailable (engine failed to load): keep short text-layer text.
        return text, ("text" if text else "error")
    try:
        bitmap = page.render(scale=OCR_SCALE).to_pil()
        buf = io.BytesIO()
        bitmap.save(buf, format="PNG")
        ocr = _ocr_text(buf.getvalue(), ocr_engine)
    except Exception:  # noqa: BLE001
        ocr = ""
    if ocr:
        return ocr, "ocr"
    return text, ("text" if text else "error")


def extractPdf(path: Path, source: str) -> Extraction:
    import pypdfium2 as pdfium

    doc = pdfium.PdfDocument(str(path))
    pages: list[PageText] = []
    try:
        for i in range(len(doc)):
            page = doc.get_page(i)
            text, method = extractPageText(page)
            pages.append(PageText(page_idx=i, text=text, method=method))
    finally:
        doc.close()
    return Extraction(source=source, path=str(path), pages=pages)


def extractImage(path: Path, source: str) -> Extraction:
    engine = _ocr_instance()
    text = _ocr_text(path.read_bytes(), engine)
    method = "ocr" if text else "error"
    return Extraction(
        source=source,
        path=str(path),
        pages=[PageText(page_idx=0, text=text, method=method)],
    )


def extractDocx(path: Path, source: str) -> Extraction:
    import docx

    document = docx.Document(str(path))

    # Paragraphs (body) plus table cells, headers, and footers. Table content
    # routinely holds lease-critical data (dates, amounts, names) that
    # paragraph-only extraction would silently drop.
    parts: list[str] = []
    try:
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
    except Exception:  # noqa: BLE001
        parts.append("")
    try:
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception:  # noqa: BLE001
        parts.append("")
    try:
        for section in document.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    parts.append(header.text.strip())
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    parts.append(footer.text.strip())
    except Exception:  # noqa: BLE001
        parts.append("")

    text = "\n".join(p for p in parts if p.strip()).strip()
    method = "paragraphs" if text else "error"
    pages = [PageText(page_idx=0, text=text, method=method)]
    return Extraction(source=source, path=str(path), pages=pages)


def extractText(path: str | Path) -> Extraction:
    """Extract text from a user document, dispatching on extension.

    Raises FileNotFoundError for missing paths, IsADirectoryError for
    directories, and ValueError for unsupported extensions.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)
    if path.is_dir():
        raise IsADirectoryError(path)
    source = path.stem
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extractPdf(path, source)
    if suffix in IMAGE_EXTS:
        return extractImage(path, source)
    if suffix in DOCX_EXTS:
        return extractDocx(path, source)
    if suffix in TEXT_EXTS or suffix == "":
        text = path.read_text(encoding="utf-8", errors="replace")
        method = "raw" if text.strip() else "error"
        return Extraction(
            source=source,
            path=str(path),
            pages=[PageText(page_idx=0, text=text, method=method)],
        )
    raise ValueError(f"unsupported file type: {suffix}")